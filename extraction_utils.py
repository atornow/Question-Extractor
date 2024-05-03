from docx import Document
import PyPDF2

# Use PyPDF to read pdf and return concatnated text
def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = []
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text.append(page.extract_text())
        return '\n'.join(text)

# Use docx to read docx, and go through any tables to extract text inside cells
def extract_text_from_docx(file_path):
    document = Document(file_path)
    tables = document.tables
    text = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text.append(paragraph.text)
    return "\n".join(text)

# Router function to chose between pdf/docx extraction
def extract_text(file_path):
    if file_path.endswith('.doc') or file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError('Unsupported file format')

# Agent that just returns 2 collumn CSV with questions text and numbers
def extract_questions(text, client, model="claude-3-sonnet-20240229"):
    message = client.messages.create(
        model=model,
        max_tokens=3000,
        temperature=0,
        system="Your task is to extract the questions from this text converted from a document into a machine-readable CSV format.\n\nThe CSV output should contain the following fields for each question:\n- question_text: The text of the question\n- question_number: The question number. For subquestions like \"b.\" under question 2.6, format as \"2.6.b\"  \n\nTo complete this task:\n\n<scratchpad>\n1. Carefully analyze the security document text to identify all questions. \n2. For each question found:\n   a. Extract the question text\n   b. Determine the question number, accounting for any subquestion structure\n3. Format each question's data into a CSV row following the specified fields\n4. Combine all question rows into a single CSV output, with one question per row\n</scratchpad>\n\nAfter completing the CSV, please output the entire CSV inside <csv> tags, like this:\n\n<csv>\nquestion_text,question_number\n\"Question 1 text\",1\n\"Question 2 text\",2\n\"Question 2 a subquestion\",2.a\n</csv>\n\nRemember, the output should always be a valid CSV that strictly matches the described format, with all questions from the input text included.",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": text
                    }
                ]
            }
        ]
    )
    return message.content[0].text if message.content else ""

# Agent that cleanse up CSV's from extract agent
def clean_csv(csv_text, client, model="claude-3-sonnet-20240229"):
    message = client.messages.create(
        model=model,
        max_tokens=3000,
        temperature=0,
        system="Your task is to clean up the provided CSV containing extracted questions. Ensure that the CSV contains only valid question records, with each field in the correct position. If any record appears malformed or has fields out of position, attempt to correct it. The CSV should maintain the following format:\n\nquestion_text,question_number\n\nAfter cleaning the CSV, please output the pure CSV nothing else.",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": csv_text
                    }
                ]
            }
        ]
    )
    return message.content[0].text if message.content else ""

# Agent that uses both raw text of doc and extracted question CSV to add 2 collumns, multiple choice boolean and answer options
def extract_answer_options(text, questions_csv, client, model="claude-3-sonnet-20240229"):
    message = client.messages.create(
        model=model,
        max_tokens=3000,
        temperature=0,
        system="Your task is to extract the answer options for the questions in the provided CSV and determine if each question is multiple choice. The CSV contains questions extracted from a text document.\n\nThe updated CSV output should contain the following fields for each question:\n- question_text: The text of the question\n- question_number: The question number. For subquestions like \"b.\" under question 2.6, format as \"2.6.b\"  \n- is_multiple_choice: 1 if the question has multiple choice options, 0 if not\n- answer_options: If is_multiple_choice is 1, a slash (/) delimited list of the possible answer options. For yes/no questions, format as \"yes/no\". Leave blank if not multiple choice.\n\nTo complete this task:\n\n<scratchpad>\n1. For each question in the provided CSV:\n   a. Carefully analyze the original text document to find the question and determine if it has answer options\n   b. If the question has answer options, set is_multiple_choice to 1 and extract a clean list of the options in the answer_options field\n   c. If the question does not have answer options, set is_multiple_choice to 0 and leave answer_options blank\n2. Output the updated CSV with the is_multiple_choice and answer_options fields added\n</scratchpad>\n\nWhen extracting answer options, be sure to robustly handle any format or style they are written in. The goal is to have a clean, complete list for each multiple choice question.\n\nAfter completing the CSV, please output the entire updated CSV inside <csv> tags, like this:\n\n<csv>\nquestion_text,question_number,is_multiple_choice,answer_options\n\"Question 1 text\",1,0,\n\"Question 2 text\",2,1,\"Option 1/Option 2/Option 3\"\n\"Question 2 a subquestion\",2.a,1,\"yes/no\"\n</csv>\n\nRemember, the output should always be a valid CSV that strictly matches the described format, with is_multiple_choice and answer options added for each question.",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"{text}\n\nExtracted Questions CSV:\n{questions_csv}"
                    }
                ]
            }
        ]
    )
    return message.content[0].text if message.content else ""

# Final pass cleaner agent that ensures 4 collumn csv with question text, number, multiple choice boolean and answer options
def clean_final_csv(csv_text, client, model="claude-3-sonnet-20240229"):
    message = client.messages.create(
        model=model,
        max_tokens=3000,
        temperature=0,
        system="Your task is to clean up the provided CSV containing extracted questions and answer options. Ensure that the CSV contains exactly 4 columns (question_text, question_number, is_multiple_choice, answer_options) in that order, with the first row being the header row. Escape any necessary characters to maintain a valid 4 column CSV structure. Output only the cleaned CSV with no extra text.",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": csv_text
                    }
                ]
            }
        ]
    )
    return message.content[0].text if message.content else ""

# Agent that uses finalized questions csv and original extracted text to create 2 collumn CSV with section desc. and section number
def extract_question_sections(text, questions_csv, client, model="claude-3-sonnet-20240229"):
    message = client.messages.create(
        model=model,
        max_tokens=3000,
        temperature=0,
        system="Your task is to extract the question section information from the provided text, using the extracted questions CSV as a reference for what not to include. Output the result as a CSV with two columns: section_number and section_description.\n\nTo complete this task:\n\n<scratchpad>\n1. Carefully analyze the text to identify sections that describe groups of questions, using the extracted questions CSV to know which specific question numbers to ignore.\n2. For each section found:\n   a. Extract the section number (e.g., 8, 8.1, etc.)\n   b. Extract the description text for that section, excluding any text that is a specific question from the CSV\n3. Format each section's data into a CSV row with two columns: section_number, section_description\n4. Combine all section rows into a single CSV output, with one section per row\n</scratchpad>\n\nAfter completing the CSV, please output only the CSV content with no extra text. The first row should be the header row with the column names.\n\nEscape any necessary characters to maintain a valid 2 column CSV structure.",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"{text}\n\nExtracted Questions CSV:\n{questions_csv}"
                    }
                ]
            }
        ]
    )
    return message.content[0].text if message.content else ""

# Cost of running API calls calculator
def calculate_cost(text, model, function_name, is_input=True):
    token_count = len(text) // 4
    
    if function_name == "extract_questions":
        system_prompt = "Your task is to extract the questions from this text converted from a document into a machine-readable CSV format.\n\nThe CSV output should contain the following fields for each question:\n- question_text: The text of the question\n- question_number: The question number. For subquestions like \"b.\" under question 2.6, format as \"2.6.b\"  \n\nTo complete this task:\n\n<scratchpad>\n1. Carefully analyze the security document text to identify all questions. \n2. For each question found:\n   a. Extract the question text\n   b. Determine the question number, accounting for any subquestion structure\n3. Format each question's data into a CSV row following the specified fields\n4. Combine all question rows into a single CSV output, with one question per row\n</scratchpad>\n\nAfter completing the CSV, please output the entire CSV inside <csv> tags, like this:\n\n<csv>\nquestion_text,question_number\n\"Question 1 text\",1\n\"Question 2 text\",2\n\"Question 2 a subquestion\",2.a\n</csv>\n\nRemember, the output should always be a valid CSV that strictly matches the described format, with all questions from the input text included."
    elif function_name == "clean_csv":
        system_prompt = "Your task is to clean up the provided CSV containing extracted questions. Ensure that the CSV contains only valid question records, with each field in the correct position. If any record appears malformed or has fields out of position, attempt to correct it. The CSV should maintain the following format:\n\nquestion_text,question_number\n\nAfter cleaning the CSV, please output the pure CSV nothing else."
    elif function_name == "extract_answer_options":
        system_prompt = "Your task is to extract the answer options for the questions in the provided CSV and determine if each question is multiple choice. The CSV contains questions extracted from a text document.\n\nThe updated CSV output should contain the following fields for each question:\n- question_text: The text of the question\n- question_number: The question number. For subquestions like \"b.\" under question 2.6, format as \"2.6.b\"  \n- is_multiple_choice: 1 if the question has multiple choice options, 0 if not\n- answer_options: If is_multiple_choice is 1, a slash (/) delimited list of the possible answer options. For yes/no questions, format as \"yes/no\". Leave blank if not multiple choice.\n\nTo complete this task:\n\n<scratchpad>\n1. For each question in the provided CSV:\n   a. Carefully analyze the original text document to find the question and determine if it has answer options\n   b. If the question has answer options, set is_multiple_choice to 1 and extract a clean list of the options in the answer_options field\n   c. If the question does not have answer options, set is_multiple_choice to 0 and leave answer_options blank\n2. Output the updated CSV with the is_multiple_choice and answer_options fields added\n</scratchpad>\n\nWhen extracting answer options, be sure to robustly handle any format or style they are written in. The goal is to have a clean, complete list for each multiple choice question.\n\nAfter completing the CSV, please output the entire updated CSV inside <csv> tags, like this:\n\n<csv>\nquestion_text,question_number,is_multiple_choice,answer_options\n\"Question 1 text\",1,0,\n\"Question 2 text\",2,1,\"Option 1/Option 2/Option 3\"\n\"Question 2 a subquestion\",2.a,1,\"yes/no\"\n</csv>\n\nRemember, the output should always be a valid CSV that strictly matches the described format, with is_multiple_choice and answer options added for each question."
    elif function_name == "clean_final_csv":
        system_prompt = "Your task is to clean up the provided CSV containing extracted questions and answer options. Ensure that the CSV contains exactly 4 columns (question_text, question_number, is_multiple_choice, answer_options) in that order, with the first row being the header row. Escape any necessary characters to maintain a valid 4 column CSV structure. Output only the cleaned CSV with no extra text."
    elif function_name == "extract_question_sections":
        system_prompt = "Your task is to extract the question section information from the provided text and output it as a CSV with two columns: section_number and section_description.\n\nTo complete this task:\n\n<scratchpad>\n1. Carefully analyze the text to identify sections that describe questions or question groups.\n2. For each section found:\n   a. Extract the section number (e.g., 8, 8.1, etc.)\n   b. Extract the description text for that section\n3. Format each section's data into a CSV row with two columns: section_number, section_description\n4. Combine all section rows into a single CSV output, with one section per row\n</scratchpad>\n\nAfter completing the CSV, please output only the CSV content with no extra text. The first row should be the header row with the column names.\n\nEscape any necessary characters to maintain a valid 2 column CSV structure."
    else:
        system_prompt = ""

    if is_input:
        token_count += len(system_prompt) // 4

    if model == "claude-3-haiku-20240307":
        cost_per_token = 0.25 / 1e6 if is_input else 1.25 / 1e6
    elif model == "claude-3-sonnet-20240229":
        cost_per_token = 3 / 1e6 if is_input else 15 / 1e6
    else:  # claude-3-opus-20240229
        cost_per_token = 15 / 1e6 if is_input else 75 / 1e6
    
    return cost_per_token * token_count