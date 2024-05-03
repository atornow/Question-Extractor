from docx import Document
import sys

def docx_to_text(file_path):
    document = Document(file_path)
    tables = document.tables
    text = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text.append(paragraph.text)
    return "\n".join(text)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python script.py <docx_file_path>")
        sys.exit(1)

    file_path = sys.argv[1]
    try:
        text = docx_to_text(file_path)
        print(text)
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)