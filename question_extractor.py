from docx import Document
import sys
import os
import anthropic
from dotenv import load_dotenv
import csv
import PyPDF2

load_dotenv(os.path.join(os.path.dirname(__file__), 'main.env'))

api_key = os.environ.get('ANTHROPIC_API_KEY')
client = anthropic.Anthropic(api_key=api_key)

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = []
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text.append(page.extract_text())
        return '\n'.join(text)

def extract_text_from_docx(file_path):
    document = Document(file_path)
    tables = document.tables
    text = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text.append(paragraph.text)
    return "\n".join(text)

def extract_text(file_path):
    if file_path.endswith('.doc') or file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError('Unsupported file format')

def clean_csv(csv_text):
    message = client.messages.create(
        model="claude-3-sonnet-20240229",
        max_tokens=3000,
        temperature=0,
        system="Your task is to clean up the provided CSV containing extracted questions. Ensure that the CSV contains only valid question records, with each field in the correct position. If any record appears malformed or has fields out of position, attempt to correct it. The CSV should maintain the following format:\n\nquestion_text,question_number\n\nAfter cleaning the CSV, please output the pure CSV nothing else.",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": csv_text
                    }
                ]
            }
        ]
    )
    return message.content[0].text if message.content else ""

def extract_questions(text):
    message = client.messages.create(
        model="claude-3-sonnet-20240229",
        max_tokens=3000,
        temperature=0,
        system="Your task is to extract the questions from this text converted from a document into a machine-readable CSV format.\n\nThe CSV output should contain the following fields for each question:\n- question_text: The text of the question\n- question_number: The question number. For subquestions like \"b.\" under question 2.6, format as \"2.6.b\"  \n\nTo complete this task:\n\n<scratchpad>\n1. Carefully analyze the security document text to identify all questions. \n2. For each question found:\n   a. Extract the question text\n   b. Determine the question number, accounting for any subquestion structure\n3. Format each question's data into a CSV row following the specified fields\n4. Combine all question rows into a single CSV output, with one question per row\n</scratchpad>\n\nAfter completing the CSV, please output the entire CSV inside <csv> tags, like this:\n\n<csv>\nquestion_text,question_number\n\"Question 1 text\",1\n\"Question 2 text\",2\n\"Question 2 a subquestion\",2.a\n</csv>\n\nRemember, the output should always be a valid CSV that strictly matches the described format, with all questions from the input text included.",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": text
                    }
                ]
            }
        ]
    )
    return message.content[0].text if message.content else ""

def extract_answer_options(text, questions_csv):
    message = client.messages.create(
        model="claude-3-sonnet-20240229",
        max_tokens=3000,
        temperature=0,
        system="Your task is to extract the answer options for the questions in the provided CSV and determine if each question is multiple choice. The CSV contains questions extracted from a text document.\n\nThe updated CSV output should contain the following fields for each question:\n- question_text: The text of the question\n- question_number: The question number. For subquestions like \"b.\" under question 2.6, format as \"2.6.b\"  \n- is_multiple_choice: 1 if the question has multiple choice options, 0 if not\n- answer_options: If is_multiple_choice is 1, a slash (/) delimited list of the possible answer options. For yes/no questions, format as \"yes/no\". Leave blank if not multiple choice.\n\nTo complete this task:\n\n<scratchpad>\n1. For each question in the provided CSV:\n   a. Carefully analyze the original text document to find the question and determine if it has answer options\n   b. If the question has answer options, set is_multiple_choice to 1 and extract a clean list of the options in the answer_options field\n   c. If the question does not have answer options, set is_multiple_choice to 0 and leave answer_options blank\n2. Output the updated CSV with the is_multiple_choice and answer_options fields added\n</scratchpad>\n\nWhen extracting answer options, be sure to robustly handle any format or style they are written in. The goal is to have a clean, complete list for each multiple choice question.\n\nAfter completing the CSV, please output the entire updated CSV inside <csv> tags, like this:\n\n<csv>\nquestion_text,question_number,is_multiple_choice,answer_options\n\"Question 1 text\",1,0,\n\"Question 2 text\",2,1,\"Option 1/Option 2/Option 3\"\n\"Question 2 a subquestion\",2.a,1,\"yes/no\"\n</csv>\n\nRemember, the output should always be a valid CSV that strictly matches the described format, with is_multiple_choice and answer options added for each question.",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"{text}\n\nExtracted Questions CSV:\n{questions_csv}"
                    }
                ]
            }
        ]
    )
    return message.content[0].text if message.content else ""

def main(file_path):
    print(f"Processing file: {file_path}")
    try:
        text = extract_text(file_path)
    except Exception as e:
        print(f"Error extracting text: {str(e)}")
        return

    try:

        questions_csv = extract_questions(text)
        print("Questions extracted:")
        print(questions_csv)

        cleaned_questions_csv = clean_csv(questions_csv)
        print("CSV cleaned:")
        print(cleaned_questions_csv)

        final_csv = extract_answer_options(text, cleaned_questions_csv)
        print("Answer options added:")
        print(final_csv)

        if final_csv:
            file_name, _ = os.path.splitext(file_path)
            csv_file_path = f"{file_name}_questions.csv"
            with open(csv_file_path, "w", newline="") as csvfile:
                csvfile.write(final_csv)
            print(f"CSV file saved as: {csv_file_path}")
        else:
            print("No content received from API, unable to save CSV file.")

    except Exception as e:
        print(f'Error: {str(e)}')

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python question_extractor.py <file_path>")
        sys.exit(1)
    file_path = sys.argv[1]
    main(file_path)